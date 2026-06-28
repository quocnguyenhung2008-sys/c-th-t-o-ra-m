from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree

# Mở rộng Namespace để quét được toàn bộ ngóc ngách của file Word
WORD_NAMESPACES = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "xml": "http://www.w3.org/XML/1998/namespace"
}


class DocxExtractionError(RuntimeError):
    pass


def extract_docx_text(path: Path) -> str:
    """
    Trích xuất toàn bộ văn bản từ file DOCX bao gồm: Thân văn bản, Bảng biểu,
    Header/Footer, và Ghi chú nhằm tối ưu hóa bộ từ khóa nhận diện môn học.
    """
    try:
        import docx  # type: ignore

        document = docx.Document(str(path))
        parts = []

        # 1. Thu thập dữ liệu từ các đoạn văn bản chính (Paragraphs)
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        # 2. Thu thập dữ liệu từ Bảng biểu (Chống trùng lặp do merged cells)
        for table in document.tables:
            seen_cells = set()  # Lưu vết các ô đã đọc để tránh lặp điểm từ khóa
            for row in table.rows:
                for cell in row.cells:
                    # Nếu ô đã được quét (do merged), bỏ qua để tránh loãng trọng số
                    if cell._tc in seen_cells:
                        continue
                    seen_cells.add(cell._tc)
                    text = cell.text.strip()
                    if text:
                        parts.append(text)

        # 3. Nâng cấp: Thu thập dữ liệu cực giá trị từ Headers và Footers
        # Đề thi THPT rất hay để tên môn học kèm năm học ở Header/Footer
        for section in document.sections:
            if section.header and not section.header.is_linked_to_previous:
                for p in section.header.paragraphs:
                    if p.text.strip():
                        parts.append(p.text.strip())
            if section.footer and not section.footer.is_linked_to_previous:
                for p in section.footer.paragraphs:
                    if p.text.strip():
                        parts.append(p.text.strip())

        return "\n".join(parts)

    except ImportError:
        # Nếu máy không cài python-docx, chuyển sang giải pháp dự phòng stdlib nâng cấp
        return _extract_docx_text_stdlib(path)
    except Exception as exc:
        raise DocxExtractionError(f"Cannot extract DOCX text from {path}: {exc}") from exc


def extract_doc_text(path: Path) -> str:
    """
    Trích xuất văn bản từ file .doc (Word 97–2003, định dạng nhị phân OLE).
    Thứ tự fallback: mammoth → antiword (CLI) → textract → lỗi rõ ràng.
    Cài nhanh: pip install mammoth
    """
    # ── Lớp 1: mammoth — thư viện Python thuần, không cần cài thêm CLI ─────
    try:
        import mammoth  # type: ignore
        with path.open("rb") as f:
            result = mammoth.extract_raw_text(f)
        text = result.value.strip()
        if text:
            return text
    except ImportError:
        pass
    except Exception as exc:
        raise DocxExtractionError(f"mammoth failed on {path}: {exc}") from exc

    # ── Lớp 2: antiword — CLI tool, cần cài riêng (brew install antiword) ──
    try:
        import subprocess
        result = subprocess.run(
            ["antiword", str(path)],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except FileNotFoundError:
        pass  # antiword chưa cài — bỏ qua, thử lớp tiếp theo
    except Exception as exc:
        raise DocxExtractionError(f"antiword failed on {path}: {exc}") from exc

    # ── Lớp 3: textract — wrapper nhiều tool, nặng hơn nhưng linh hoạt ─────
    try:
        import textract  # type: ignore
        text = textract.process(str(path)).decode("utf-8", errors="ignore").strip()
        if text:
            return text
    except ImportError:
        pass
    except Exception as exc:
        raise DocxExtractionError(f"textract failed on {path}: {exc}") from exc

    raise DocxExtractionError(
        f"Không thể đọc file .doc: {path.name}. "
        "Cài mammoth để hỗ trợ: pip install mammoth"
    )


def _extract_docx_text_stdlib(path: Path) -> str:
    """
    Giải pháp dự phòng bằng thư viện chuẩn: Quét sâu vào cấu trúc XML của gói Zip.
    Khắc phục triệt để lỗi dính chữ và mất thuộc tính giữ khoảng trắng (xml:space="preserve").
    """
    parts = []
    # Các file XML chứa nội dung tiềm năng trong gói cấu trúc file Word
    target_xmls = [
        "word/document.xml",
        "word/footnotes.xml",
        "word/endnotes.xml",
        "word/header1.xml", "word/header2.xml", "word/header3.xml",
        "word/footer1.xml", "word/footer2.xml", "word/footer3.xml"
    ]

    try:
        with zipfile.ZipFile(path) as archive:
            available_files = set(archive.namelist())

            for xml_file in target_xmls:
                if xml_file not in available_files:
                    continue

                xml_content = archive.read(xml_file)
                root = ElementTree.fromstring(xml_content)

                # Duyệt qua các thẻ đoạn văn <w:p> thay vì lấy thô thẻ <w:t> để kiểm soát khoảng trắng
                for p_node in root.findall(".//w:p", WORD_NAMESPACES):
                    p_texts = []
                    for t_node in p_node.findall(".//w:t", WORD_NAMESPACES):
                        # GIẢI QUYẾT LỖI DÍNH CHỮ: Kiểm tra thuộc tính bảo toàn khoảng trắng xml:space="preserve"
                        space_attr = t_node.get(f"{{{WORD_NAMESPACES['xml']}}}space")
                        text = t_node.text or ""

                        if space_attr == "preserve":
                            p_texts.append(text)
                        else:
                            p_texts.append(text.strip())

                    # Gộp các cụm từ trong cùng 1 đoạn bằng khoảng trắng
                    paragraph_str = " ".join(t for t in p_texts if t).strip()
                    if paragraph_str:
                        parts.append(paragraph_str)

    except Exception as exc:
        if "word/document.xml" not in parts:
            raise DocxExtractionError(f"Cannot read DOCX package {path}: {exc}") from exc

    return "\n".join(parts)